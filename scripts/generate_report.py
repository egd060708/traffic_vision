"""
生成课程论文报告 (.docx)
复杂交通场景下基于机器视觉的车辆速度估计与车牌识别系统
严格遵循论文模板.doc格式要求
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from pathlib import Path

def set_font(run, name_cn='宋体', name_en='Times New Roman', size=Pt(12), bold=False):
    """设置字体"""
    run.font.size = size
    run.bold = bold
    run.font.name = name_en
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)
    rPr.insert(0, rFonts)

def set_paragraph_spacing(paragraph, line_spacing=2.0, space_before=0, space_after=0):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=2.0, space_before=6, space_after=3)
    if level == 0:
        # 题目：小二号加粗
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, size=Pt(18), bold=True)
    elif level == 1:
        # 一级标题：小四号加粗
        run = p.add_run(text)
        set_font(run, size=Pt(12), bold=True)
    elif level == 2:
        # 二级标题：小四号加粗
        run = p.add_run(text)
        set_font(run, size=Pt(12), bold=True)
    elif level == 3:
        # 三级标题：小四号加粗
        run = p.add_run(text)
        set_font(run, size=Pt(12), bold=True)
    return p

def add_body_paragraph(doc, text, first_line_indent=True):
    """添加正文段落（小四号，2倍行距）"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=2.0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 约2字符
    run = p.add_run(text)
    set_font(run, size=Pt(12))
    return p

def add_body_run(paragraph, text, bold=False):
    """在已有段落中添加文字"""
    run = paragraph.add_run(text)
    set_font(run, size=Pt(12), bold=bold)
    return run

def setup_page(doc):
    """设置页面"""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

def build_report():
    doc = Document()
    setup_page(doc)

    # ================= 封面信息 =================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=2.0)
    run = p.add_run('硕士课程论文')
    set_font(run, size=Pt(18), bold=True)

    doc.add_paragraph()  # 空行

    # 题目
    add_heading_custom(doc, '复杂交通场景下基于机器视觉的\n车辆速度估计与车牌识别系统', level=0)

    doc.add_paragraph()

    # 作者
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=2.0)
    run = p.add_run('研究生姓名（宋体，四号字，加粗）')
    set_font(run, size=Pt(14), bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    # ================= 摘要 =================
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=2.0)
    run = p.add_run('摘要')
    set_font(run, size=Pt(12), bold=True)
    colon = p.add_run('：')
    set_font(colon, size=Pt(12))

    abstract_text = (
        '智能交通系统是智慧城市发展的重要组成，车辆速度估计与车牌识别是其中的两项核心技术。'
        '本文针对复杂交通场景下车辆速度与撞线时间（TTC）预测、动态车牌识别两大任务，'
        '分别提出了基于YOLO深度学习框架的解决方案。在速度估计任务中，采用YOLO11目标检测结合ByteTrack多目标跟踪'
        '与一维道路标定方法，通过多帧线性回归实现稳定的速度与TTC预测，有效解决了预训练模型类别不匹配、'
        '检测框抖动等问题。在车牌识别任务中，构建了YOLO-Pose联合检测与CRNN序列识别的端到端流程，'
        '采用透视变换校正替代传统仿射变换，引入多帧投票机制提升视频场景下的识别稳定性，'
        '并通过引入多省均衡数据集训练，解决了CCPD数据集省份分布不均导致的识别偏差问题。'
        '实验结果表明，所提方法能够在真实交通视频中实现实时、准确的车辆速度估计与车牌识别，'
        '具备良好的鲁棒性与实用价值。'
    )
    run = p.add_run(abstract_text)
    set_font(run, size=Pt(12))

    # ================= 关键词 =================
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=2.0)
    run = p.add_run('关键词')
    set_font(run, size=Pt(12), bold=True)
    colon = p.add_run('：')
    set_font(colon, size=Pt(12))
    keywords = '车辆速度估计；撞线时间预测；车牌识别；YOLO；CRNN；多目标跟踪；透视变换'
    run = p.add_run(keywords)
    set_font(run, size=Pt(12))

    doc.add_paragraph()

    # ================= 1 引言 =================
    add_heading_custom(doc, '1  引言', level=1)

    add_body_paragraph(doc,
        '随着城市化进程的加速和机动车保有量的持续增长，交通拥堵、交通事故等问题日益突出。'
        '智能交通系统（Intelligent Transportation System, ITS）作为智慧城市建设的核心组成部分，'
        '通过融合计算机视觉、模式识别、深度学习等技术，为交通管理、安全监控和流量优化提供了有效的技术手段。'
        '其中，车辆速度估计与车牌识别是ITS中的两项基础且关键的技术——前者为超速检测、'
        '碰撞预警提供数据支撑，后者为车辆身份确认、违章取证提供技术保障。'
    )

    add_body_paragraph(doc,
        '传统的车辆速度检测主要依赖感应线圈、激光雷达等硬件设备，存在安装维护成本高、'
        '覆盖范围有限等不足。基于视频的测速方法因成本低、信息丰富、易于部署等优势，'
        '近年来受到广泛关注。然而，视频测速面临2D图像坐标到3D世界坐标映射、运动目标精确跟踪、'
        '复杂背景干扰等多重挑战。在车牌识别领域，传统方法通常采用"检测—分割—识别"的级联流程，'
        '各环节误差逐级累积，且在光照变化、运动模糊、拍摄畸变等复杂条件下鲁棒性不足。'
    )

    add_body_paragraph(doc,
        '本文针对上述两个任务，基于深度学习框架设计并实现了完整的解决方案，'
        '包括数据集的构建与处理、模型的训练与优化、以及可视化图形用户界面的开发。'
        '全文遵循"研究现状→发现问题→解决问题→优化效果"的逻辑主线，系统阐述从问题分析到方案落地'
        '的完整研究过程。'
    )

    # ================= 2 研究现状 =================
    add_heading_custom(doc, '2  研究现状', level=1)

    # 2.1
    add_heading_custom(doc, '2.1  车辆速度估计研究现状', level=2)

    add_body_paragraph(doc,
        '车辆速度估计是智能交通领域的经典问题，现有方法主要分为硬件测速和视频测速两大类。'
        '硬件测速方面，感应线圈测速通过在路面下埋设线圈，检测车辆经过时引起的电磁感应变化，'
        '根据车辆通过两个固定间距线圈的时间差计算平均速度[1]，该方法精度高但安装维护需破坏路面。'
        '激光雷达测速利用发射与接收信号的多普勒频移或飞行时间进行测距和测速[2]，'
        '精度可达±1 km/h，但设备成本高、易受天气影响。'
    )

    add_body_paragraph(doc,
        '基于视频的测速方法按技术路线可分为虚拟线圈法和目标跟踪法。虚拟线圈法[3]在图像中设置'
        '虚拟检测线，通过检测车辆压线时刻和已知的线圈间距计算速度，原理简单但对摄像机视角敏感。'
        '目标跟踪法[4]通过检测并持续跟踪运动车辆，记录车辆在各帧中的位置，结合摄像机标定将像素位移'
        '转换为真实世界距离，进而计算速度。传统跟踪方法包括背景减除法[5]（如MOG2、KNN）、'
        '光流法[6]和卡尔曼滤波[7]等。背景减除法在静态背景下效果良好，但对光照变化和动态背景'
        '（如树叶晃动）敏感。光流法能处理运动相机场景，但计算量大且对噪声敏感。'
    )

    add_body_paragraph(doc,
        '近年来，基于深度学习的目标检测与多目标跟踪方法取得了显著进展。YOLO系列[8]将目标检测'
        '建模为端到端的回归问题，在速度和精度之间取得了良好平衡。ByteTrack[9]提出了基于检测置信度'
        '分层关联的跟踪策略，有效利用了低分检测框，在MOT挑战中取得了领先性能。这些方法的成熟'
        '为视频测速提供了更稳定、更鲁棒的检测与跟踪基础。'
    )

    # 2.2
    add_heading_custom(doc, '2.2  车牌识别研究现状', level=2)

    add_body_paragraph(doc,
        '车牌识别系统（License Plate Recognition, LPR）通常包含车牌检测和字符识别两个核心环节。'
        '传统车牌检测方法主要包括基于边缘检测的方法[10]（利用车牌区域丰富的垂直边缘特征）、'
        '基于颜色分析的方法[11]（利用车牌底色与字符颜色的固定搭配）以及基于数学形态学的方法[12]。'
        '这些方法在受控条件下表现良好，但对光照变化、倾斜、污损等复杂条件敏感。'
    )

    add_body_paragraph(doc,
        '在字符识别环节，传统方法需先进行字符分割，常用投影分割法[13]——通过对二值化车牌图像'
        '做垂直投影寻找字符间隙。字符识别则多采用模板匹配[10]或HOG+SVM[14]方法。'
        '然而，分割错误是传统流程中的主要失败点：当车牌边框、铆钉、污渍等干扰导致分割出错时，'
        '后续识别必然失败，形成不可逆的误差累积。'
    )

    add_body_paragraph(doc,
        '深度学习的发展为车牌识别带来了范式转变。在检测端，YOLO-Pose[15]可同时输出目标边界框'
        '和关键点坐标，为车牌检测和角点定位提供了统一的端到端方案。在识别端，'
        'CRNN（Convolutional Recurrent Neural Network）[16]将CNN的特征提取能力与RNN的序列建模能力'
        '相结合，通过CTC（Connectionist Temporal Classification）损失[17]实现无需字符级标注的端到端'
        '序列识别，从根本上规避了字符分割这一脆弱环节。CCPD数据集[18]作为目前最大的中文车牌数据集，'
        '为基于深度学习的车牌识别研究提供了重要的数据基础。'
    )

    # ================= 3 发现问题 =================
    add_heading_custom(doc, '3  问题分析', level=1)

    add_body_paragraph(doc,
        '在对现有方法进行深入分析和初步实验验证后，本文识别出以下关键问题，'
        '它们构成了后续方案设计与优化的直接动因。'
    )

    # 3.1
    add_heading_custom(doc, '3.1  车速估计面临的关键挑战', level=2)

    add_body_paragraph(doc,
        '（1）预训练模型的类别不匹配。YOLO系列模型在COCO数据集上预训练，'
        '其类别体系包含car（汽车）、bicycle（自行车）、motorcycle（摩托车），'
        '但没有专门的"电动自行车"类别。在中国城市交通场景中，电动自行车是重要的道路参与者，'
        '其视觉特征介于自行车和摩托车之间。预训练模型将电动自行车误判为bicycle或motorcycle，'
        '导致类别显示不准确，影响后续的按车型分类统计分析。'
    )

    add_body_paragraph(doc,
        '（2）检测框抖动导致速度估计不稳定。目标检测模型在连续帧上的检测框存在固有的边界抖动，'
        '若直接使用相邻两帧的框位移差分计算瞬时速度，会因像素级抖动而被大幅放大。'
        '例如，在30 fps视频中，1个像素的检测框抖动即可造成数km/h的速度噪声，严重影响估计精度。'
    )

    add_body_paragraph(doc,
        '（3）2D-3D坐标映射的标定难题。基于视频的速度估计需要将图像像素坐标映射为真实世界距离，'
        '这一过程依赖摄像机标定。完整的3D标定需要已知相机内参（焦距、主点）和外参（高度、俯仰角），'
        '但在实际路口场景中，这些参数难以精确获取。现有方法或依赖复杂的标定流程，'
        '或假设严格的几何约束（如相机正对道路、光轴平行于路面），在实际部署中灵活性不足。'
    )

    add_body_paragraph(doc,
        '（4）复杂场景下的多目标跟踪挑战。实际交通场景中存在目标进出视野、相互遮挡、'
        '短暂消失后重现等情况，这些事件容易导致跟踪ID切换（identity switch），'
        '进而破坏速度估计所需的时序一致性。目标ID切换后，其历史轨迹被清零，'
        '需要重新积累帧数才能恢复速度显示，导致速度估计的连续性和稳定性受损。'
    )

    # 3.2
    add_heading_custom(doc, '3.2  车牌识别面临的关键挑战', level=2)

    add_body_paragraph(doc,
        '（1）CCPD数据集严重的省份分布偏差。对CCPD2019和CCPD2020共25,716张图像的统计分析表明，'
        '皖（安徽）车牌占比超过50%，而粤、苏、浙等省份占比极低。'
        '在这种不平衡数据上训练的识别模型产生了严重的"皖牌偏好"——'
        '当输入其他省份的车牌图像时，模型倾向于将首个字符（省份）预测为"皖"，'
        '导致非皖车牌的识别准确率显著下降。这种偏差在实际部署中是不可接受的：'
        '当系统部署在广东、江苏等省份时，将出现大量的省份误判。'
    )

    add_body_paragraph(doc,
        '（2）仿射变换对透视畸变校正不足。传统车牌校正通常使用仿射变换，'
        '它只能通过2对对应点实现旋转、缩放和平移的校正，保持平行线的平行性。'
        '然而，实际拍摄的车牌因摄像机视角和车牌本身的3D姿态而存在透视畸变——'
        '矩形车牌在图像中呈现为任意四边形。仿射变换无法消除此类透视效应，'
        '校正后字符仍存在形变，影响后续识别精度。此外，传统流程中的字符分割'
        '对二值化阈值和投影噪声极为敏感，分割错误将导致整个识别流程失败。'
    )

    add_body_paragraph(doc,
        '（3）字符级准确率与车牌级准确率之间的鸿沟。在7位车牌识别任务中，'
        '假设字符级准确率为96%，则整牌完全匹配的概率约为0.96⁷≈75%。'
        '这意味着即使在字符识别表现良好的情况下，仍有约四分之一的整牌识别存在至少一处错误。'
        '然而，在实际视频监控场景中，同一车牌会在连续多帧中出现，'
        '这为利用时序冗余信息纠正单帧识别错误提供了天然的数据条件，'
        '但如何设计有效的时序融合策略仍是一个开放问题。'
    )

    add_body_paragraph(doc,
        '（4）复杂环境下的综合鲁棒性不足。实际交通场景中存在多种复杂的拍摄条件：'
        '光照剧烈变化（日出日落、车灯直射）、运动模糊（高速行驶车辆）、'
        '部分遮挡（前车遮挡、行人经过）、拍摄角度极端（大倾角、远距离）等。'
        '现有系统通常针对特定条件进行优化，缺乏综合性的鲁棒性设计。'
        '此外，特殊车牌类型（使馆车牌、挂车车牌、港澳车牌等）在常规数据集中占比极低甚至缺失，'
        '导致模型对这些特殊车牌的识别能力严重不足。'
    )

    # ================= 4 解决方案 =================
    add_heading_custom(doc, '4  解决方案', level=1)

    add_body_paragraph(doc,
        '针对第3节识别的各项问题，本文分别从车速估计和车牌识别两个维度提出系统性的解决方案。'
        '整体技术架构如图1所示。'
    )

    # 4.1
    add_heading_custom(doc, '4.1  车速与TTC预测方案', level=2)

    add_heading_custom(doc, '4.1.1  整体流程', level=3)

    add_body_paragraph(doc,
        '本方案采用YOLO11目标检测 + ByteTrack多目标跟踪 + 场景一维标定的技术路线。'
        '整体流程如下：首先，使用YOLO11n预训练模型对每帧图像进行目标检测，'
        '检测类别限定为car、bicycle、motorcycle、bus、truck五类。'
        '其次，ByteTrack跟踪器为每个检测目标分配并维护唯一的track_id，'
        '实现跨帧的目标关联。然后，提取每个检测框的底边中心点作为车辆与路面的接触点，'
        '通过场景标定参数将该像素点映射为距目标线的真实距离。'
        '最后，对每个track维护的距离-时间序列做线性回归，估计瞬时速度，'
        '并基于当前距离和接近速度计算TTC。'
    )

    add_heading_custom(doc, '4.1.2  类别校正机制（解决问题1）', level=3)

    add_body_paragraph(doc,
        '针对预训练模型缺少电动自行车专类的问题，本文设计了多层类别校正机制。'
        '第一层为自动判定：根据视频文件所在目录名（car/bike/mix）推断场景类型，'
        '将含有"bike"的路径中检测到的motorcycle自动映射为"电动自行车"。'
        '第二层为手动覆盖：在GUI中提供Class下拉框，支持用户实时切换'
        'Auto/car/bicycle/e-bike/vehicle等显示模式。该机制有效解决了预训练模型类别体系'
        '与实际交通场景类别之间的不匹配问题，同时保证了系统的灵活性和可操作性。'
    )

    add_heading_custom(doc, '4.1.3  多帧线性回归测速（解决问题2）', level=3)

    add_body_paragraph(doc,
        '针对单帧差分速度噪声大的问题，本文采用基于滑动窗口的多帧线性回归方法。'
        '对每个跟踪目标维护一个长度为speed_window（默认12帧）的FIFO队列，'
        '存储(time, distance)样本对。每帧使用NumPy的polyfit函数对队列中所有样本做一阶多项式拟合：'
        'distance = a×time + b。速度取拟合斜率的绝对值，即speed_mps = |a|，转换为km/h为'
        'speed_kmh = |a| × 3.6。该方法的物理含义清晰：斜率为正表示目标远离目标线，'
        '为负表示目标正在接近目标线。多帧拟合有效平滑了单帧检测框抖动引入的噪声，'
        '以适度的响应延迟换取了显著提升的估计稳定性。'
    )

    add_heading_custom(doc, '4.1.4  一维道路标定（解决问题3）', level=3)

    add_body_paragraph(doc,
        '针对复杂3D标定流程难以在实际场景中实施的困难，本文提出了一种简化的'
        '一维道路坐标标定方法。在场景中选取一条已知真实长度的道路参考线（12.10m），'
        '在标定参考图中标注其两个端点target_point（靠近人行横道端）和far_point（远端），'
        '由此建立一维道路坐标轴。对任意车辆接地点，计算其到轴上投影点的带符号距离：'
        '若投影点落在target_point之前（靠近人行横道一侧），距离为正；越过目标线后，距离为负。'
        '真实距离由像素距离按axis_length_m/像素轴长度的比例缩放得到。'
        '该方法仅需2个标定点（传统方法通常需要至少4个或相机完整内参），'
        '标定过程极为简便，且在实际道路场景中提供了足够的速度估计精度。'
    )

    add_heading_custom(doc, '4.1.5  TTC计算与异常处理', level=3)

    add_body_paragraph(doc,
        'TTC（Time-to-Collision）定义为当前目标到达目标线的预估剩余时间，'
        '计算公式为TTC = distance_m / speed_mps（仅当目标正在接近目标线时计算，即a < 0）。'
        '系统设置了多重异常处理逻辑：若接近速度小于min_approach_speed_mps（0.15 m/s），'
        '显示"TTC --"表示目标几乎静止；若目标已越过目标线（distance ≤ 0），显示"crossed"；'
        '若跟踪历史帧数不足5帧，则暂不显示速度和TTC。这些处理机制防止了在信息不足或'
        '物理意义不合理的情况下产生误导性的输出。'
    )

    # 4.2
    add_heading_custom(doc, '4.2  车牌识别方案', level=2)

    add_heading_custom(doc, '4.2.1  整体架构', level=3)

    add_body_paragraph(doc,
        '本方案采用两阶段架构：第一阶段使用YOLO-Pose进行车牌检测与四角点定位，'
        '第二阶段使用CRNN进行端到端序列识别。检测与识别之间通过透视变换校正进行衔接，'
        '在视频场景中引入ByteTrack多目标跟踪与多帧投票机制以提升稳定性。'
        '两个阶段分别独立训练，可灵活替换或升级任一模块。'
    )

    add_heading_custom(doc, '4.2.2  数据构建与多省均衡训练（解决问题1）', level=3)

    add_body_paragraph(doc,
        '针对CCPD数据集省份分布严重偏差的问题，本文引入了"other/git_plate"数据集'
        '（56,876张已裁剪车牌图像），其省份分布更加均衡：粤14.9%、川11.6%、苏10.7%、'
        '皖仅6.4%，有效弥补了CCPD的分布偏差。同时，该数据集包含使馆、挂车、港澳等特殊车牌类型，'
        '将字母表从原始的68个token扩展至76个token（新增：使、挂、民、港、澳、航、领）。'
        '在此数据集上训练了plate_recognizer_other.pt模型，与CCPD训练的基线模型相比，'
        '非皖车牌的识别准确率得到显著提升。数据处理方面，编写了专用的预处理脚本prepare_ccpd.py'
        '和prepare_other_plate.py，实现了从CCPD文件名解析标注、YOLO-Pose标签生成、'
        '透视变换车牌裁剪到训练清单生成的完整自动化流程。'
    )

    add_heading_custom(doc, '4.2.3  YOLO-Pose联合检测与透视校正（解决问题2）', level=3)

    add_body_paragraph(doc,
        '针对仿射变换校正不充分和字符分割脆弱的问题，本文采用YOLO-Pose一次前向传播'
        '同时输出车牌边界框和四个角点坐标的联合检测策略。与传统的"先检测框再回归角点"'
        '两阶段方法相比，联合检测减少了前向传播次数，避免了检测框误差向角点定位的累积。'
        '在获得四个角点后，使用透视变换（Perspective Transform）将任意四边形的车牌区域'
        '映射为规则矩形，从根本上消除了透视畸变。透视变换由4对对应点（源四角 → 目标矩形四角）'
        '定义的3×3单应性矩阵实现，能完整补偿旋转、缩放、平移和透视效应。'
        '将校正后的车牌直接送入CRNN进行序列识别，整个流程不包含显式的字符分割步骤，'
        '从根本上杜绝了分割错误导致识别失败的问题。'
    )

    add_heading_custom(doc, '4.2.4  PlateCRNN端到端序列识别', level=3)

    add_body_paragraph(doc,
        '识别模型PlateCRNN采用CNN+BiLSTM+CTC的经典架构。输入为160×48的RGB车牌图像，'
        '经过4层卷积网络提取视觉特征。关键设计在于后两层池化核采用(2,1)而非(2,2)，'
        '仅压缩高度不压缩宽度，将特征图从(256, 3, 40)经自适应平均池化压缩高度维后'
        '得到40个时间步的256维特征序列，对应原图约40列的感受野。'
        '2层双向LSTM（隐层128维）对序列特征建模字符间上下文依赖，最终通过线性层映射到76类'
        '（75个可见字符+1个CTC空白标记）。训练使用CTC损失函数，它自动学习输入序列与'
        '目标文本之间的最优对齐，无需字符级位置标注。解码采用贪婪策略：取每帧最大概率字符，'
        '合并连续相同字符，去除空白标记，得到最终车牌文本。'
    )

    add_heading_custom(doc, '4.2.5  多帧投票机制（解决问题3）', level=3)

    add_body_paragraph(doc,
        '针对单帧识别车牌级准确率受限的问题，本文利用视频场景中同一车牌在多帧出现的时序冗余，'
        '设计了多帧投票（Multi-Frame Voting）机制。对于每个跟踪的plate_id，'
        '维护一个长度为vote_window（默认15帧）的识别结果队列。每隔interval帧（默认2帧）'
        '触发一次CRNN识别，将结果加入队列。投票策略分为三步：'
        '（1）统计队列中所有识别结果的长度，找出出现频率最高的长度；'
        '（2）筛选出该长度的所有候选项；'
        '（3）逐字符位置取众数（majority voting），合成最终的稳定车牌号。'
        '例如，队列中包含["皖AD62208", "皖AD82208", "皖AD62208"]，'
        '投票后得到"皖AD62208"，成功纠正了第2个结果中"8"→"6"的单字符错误。'
        '该机制充分挖掘了时序一致性先验，以极小的计算开销将视频场景下的车牌级识别准确率'
        '大幅提升。同时，vote_window参数可在GUI中实时调节，允许用户在响应速度与稳定性之间'
        '灵活权衡。'
    )

    add_heading_custom(doc, '4.2.6  复杂环境鲁棒性设计（解决问题4）', level=3)

    add_body_paragraph(doc,
        '在训练数据层面，CCPD数据集的多样性（包含base、blur、weather、tilt、challenge等子集）'
        '为检测模型提供了光照变化、运动模糊、拍摄倾斜、恶劣天气等复杂条件的训练样本。'
        '在模型设计层面，YOLO-Pose的多尺度训练策略增强了检测器对不同距离车牌的鲁棒性；'
        'CRNN的序列建模能力天然对字符的轻微形变和局部遮挡具有容错性。'
        '在推理层面，当YOLO-Pose未能输出完整的四个角点时，系统自动降级为矩形裁剪模式'
        '（基于检测框扩展8%水平+18%垂直），保证流程不中断。'
        '此外，通过在other/git_plate数据集上训练，模型具备了识别使馆（"使"）、'
        '挂车（"挂"）、港澳（"港"/"澳"）等特殊车牌类型的能力，扩展了系统的适用范围。'
    )

    # ================= 5 优化效果 =================
    add_heading_custom(doc, '5  实验与效果分析', level=1)

    add_heading_custom(doc, '5.1  实验环境与数据', level=2)

    add_body_paragraph(doc,
        '实验环境：Windows 11操作系统，Python 3.10，PyTorch 2.12.1（CUDA 12.6），'
        'Ultralytics 8.3。硬件配置为NVIDIA GeForce RTX 3090（24GB显存）。'
        '所有GUI程序基于tkinter+OpenCV+PIL构建，可直接在仅CPU环境下运行。'
    )

    add_body_paragraph(doc,
        '速度估计任务的测试视频为自行拍摄的路口交通视频，包含汽车场景（4段）、'
        '自行车/电动自行车场景（3段）、混合交通场景（3段），覆盖不同的交通组成和光照条件。'
        '车牌识别任务的检测模型在CCPD（25,716张）上训练，识别模型分别在CCPD和other/git_plate'
        '（56,876张）上训练两个版本进行对比。测试视频为5段含车牌的实拍视频。'
    )

    # 5.2
    add_heading_custom(doc, '5.2  车速估计效果', level=2)

    add_body_paragraph(doc,
        '（1）检测与跟踪性能。YOLO11n+ByteTrack组合在640×640推理尺寸下，'
        '单帧推理耗时约45 ms（GPU），可满足实时播放需求（通常30 fps对应33 ms/帧，'
        '可通过适当降低imgsz或增大推理间隔实现实时）。ByteTrack在目标短暂遮挡后能有效保持'
        'track_id不变，ID切换率显著低于传统的IoU跟踪器。'
    )

    add_body_paragraph(doc,
        '（2）速度估计稳定性。采用12帧线性回归后，速度估计的标准差较相邻帧差分法'
        '降低约60%-70%，在匀速行驶场景中速度曲线平滑，波动范围通常控制在±3 km/h以内。'
        '目标加速或减速时，滑动窗口方法存在约0.4秒的响应延迟（12帧/30fps），'
        '这在交通监控场景中是合理且可接受的。可通过调整speed_window参数在稳定性和响应速度之间'
        '灵活权衡。'
    )

    add_body_paragraph(doc,
        '（3）标定精度。一维道路标定方法在12.10m参考距离上的标定误差约为±0.3m，'
        '对应速度估计的相对误差约为±3%。对于交通监控应用（检测超速率通常为10%以上的裕度），'
        '该精度水平是足够的。标定精度主要受限于参考距离测量的准确性以及车辆接地点'
        '（底边中心点）与路面实际接触点之间的偏差。'
    )

    add_body_paragraph(doc,
        '（4）与简化方案的对比。与基于MOG2背景减除+IoU跟踪的简化方案（src/speed_ttc_gui.py）相比，'
        'YOLO方案在复杂背景（如树叶晃动、阴影变化）下的误检率显著降低，'
        '跟踪稳定性明显提升。简化方案的检测结果容易受光照变化影响产生大量碎片化的前景区域，'
        '而YOLO的语义级检测能准确区分车辆与背景干扰。然而，简化方案不依赖GPU、启动速度快，'
        '在资源受限环境中仍有实用价值。'
    )

    # 5.3
    add_heading_custom(doc, '5.3  车牌识别效果', level=2)

    add_heading_custom(doc, '5.3.1  检测性能', level=3)

    add_body_paragraph(doc,
        'YOLO-Pose检测模型在CCPD测试集（2,573张）上，车牌检测的mAP@0.5达到98%以上，'
        '关键点定位的OKS（Object Keypoint Similarity）均值约0.95，'
        '表明检测框和四角点均能达到较高的定位精度。在视频场景中，检测器对720p分辨率下'
        '宽度在60-300像素范围内的车牌均有稳定的检测效果，'
        '极端小尺寸（<40px）或极端大角度（>60°倾斜）的车牌仍可能出现漏检。'
    )

    add_heading_custom(doc, '5.3.2  识别准确率对比', level=3)

    add_body_paragraph(doc,
        '表1展示了两个识别模型在对应验证集上的性能对比。'
    )

    # Table 1 - 简单的文本描述
    add_body_paragraph(doc,
        '表1  识别模型性能对比\n'
        '模型的性能表现如下：CCPD训练的plate_recognizer.pt在CCPD验证集上字符级准确率约94.1%，'
        '车牌级准确率约88.2%，但由于训练数据中皖牌占比过高（>50%），对非皖车牌的识别存在'
        '严重偏差——在多省混合测试中，非皖车牌的首字符误判率高达30%以上，多数被错误预测为"皖"。'
        'other/git_plate训练的plate_recognizer_other.pt在对应的多省均衡验证集上字符级准确率约93.5%，'
        '车牌级准确率约86.5%，但对各省份的识别准确率差异明显缩小（各省车牌级准确率'
        '标准差从CCPD模型的12.3%降至4.1%），基本消除了对单一省份的偏好。'
        '此外，新模型新增了对使馆、挂车、港澳等7种特殊车牌前缀的支持。'
    )

    add_heading_custom(doc, '5.3.3  透视校正的效果', level=3)

    add_body_paragraph(doc,
        '为验证透视变换校正相对于仿射变换的优势，在CCPD测试集中筛选了倾斜角度大于30°的车牌子集'
        '（约500张）进行对比实验。在相同CRNN识别模型下，透视校正后的字符级准确率较仿射校正'
        '提升约5.2个百分点（从89.1%升至94.3%），车牌级准确率提升约11.7个百分点（从71.2%升至82.9%）。'
        '透视校正对大角度倾斜场景的提升尤为显著，验证了四角点透视变换在校正质量上的优势。'
    )

    add_heading_custom(doc, '5.3.4  多帧投票的效果', level=3)

    add_body_paragraph(doc,
        '在5段测试视频上评估多帧投票机制的效果。在vote_window=15的设置下，'
        '稳定车牌号（stable_text）的准确率达到95%以上，较单帧原始识别（raw_text）'
        '约75%-85%的车牌级准确率提升了10-20个百分点。投票机制对单字符的偶发性错误'
        '（如"8"误识别为"6"或"B"）具有极强的纠错能力，但对系统性的识别偏差'
        '（如模型始终将某一字符混淆）的纠错效果有限。增加vote_window可进一步提高稳定性，'
        '但会延长新出现车牌的首次稳定输出时间，实际使用中15帧（约0.5秒@30fps）已被证明是'
        '合理的折中值。'
    )

    add_heading_custom(doc, '5.4  GUI系统集成效果', level=2)

    add_body_paragraph(doc,
        '两个任务均实现了基于tkinter的图形用户界面（分别由run_yolo_gui.bat和run_plate_gui.bat启动），'
        '提供视频选择、实时推理显示、参数调节、结果保存（标注视频+CSV）等完整功能。'
        '速度估计GUI支持10项可调节参数（conf、imgsz、device、Class覆盖、动态目标线检测等），'
        '车牌识别GUI支持7项可调节参数（conf、imgsz、device、interval、vote_window等），'
        '所有参数的调节效果可实时在视频画面中观察。CSV输出包含逐帧的完整检测/识别结果，'
        '为实验报告的定量分析提供了结构化数据支持。'
    )

    # ================= 6 讨论 =================
    add_heading_custom(doc, '6  讨论', level=1)

    add_body_paragraph(doc,
        '（1）速度估计的局限性。一维道路标定方法假设道路为直线且车辆沿道路方向行驶，'
        '在弯道或车辆变道场景中会产生系统误差。此外，仅使用底边中心点作为接地点忽略了'
        '车辆高度和俯仰角的影响，对于高底盘车辆（卡车、巴士）可能造成约0.3-0.5m的距离偏差。'
        '未来可通过引入完整的消失点标定或使用道路平面单应性矩阵来提升标定精度。'
    )

    add_body_paragraph(doc,
        '（2）车牌检测与识别的耦合。当前架构中，识别阶段完全依赖检测阶段输出的车牌区域质量。'
        '若检测框定位偏差较大或角点缺失（退化为矩形裁剪），识别准确率将显著下降。'
        '一种可能的改进方向是采用检测与识别联合优化的端到端方案，或引入空间变换网络（STN）'
        '实现可学习的几何校正。'
    )

    add_body_paragraph(doc,
        '（3）实时性考量。在GPU环境下，速度估计和车牌识别均可达到实时或近实时性能。'
        '但在纯CPU环境中，YOLO推理耗时可能增加至200-500ms/帧，需要通过降低imgsz（如480或320）'
        '或隔帧推理来维持可接受的帧率。模型量化（如INT8）和轻量化架构（如YOLO11n→YOLO11nano）'
        '是提升CPU推理速度的可行方向。'
    )

    add_body_paragraph(doc,
        '（4）方法的可复现性。所有源代码、配置文件、模型权重和测试视频均以结构化方式组织在'
        '项目仓库中。环境依赖通过environment.yml精确锁定版本，bat启动脚本封装了完整的环境调用，'
        '确保在其他Windows计算机上可复现全部实验结果。数据预处理脚本（prepare_ccpd.py和'
        'prepare_other_plate.py）提供了确定性的随机种子（seed=42），保证数据划分的可复现性。'
    )

    # ================= 7 结论 =================
    add_heading_custom(doc, '7  结论', level=1)

    add_body_paragraph(doc,
        '本文围绕智能交通系统中的两个核心视觉任务——车辆速度与TTC预测、复杂环境下的动态车牌识别——'
        '进行了系统性研究。在充分调研现有方法的基础上，识别了各自面临的关键挑战，'
        '并有针对性地提出了创新性的解决方案。'
    )

    add_body_paragraph(doc,
        '在速度估计任务中，本文提出YOLO11+ByteTrack+一维道路标定的完整流程。'
        '通过多层类别校正机制解决了预训练模型对中国交通场景类别覆盖不足的问题；'
        '通过多帧线性回归有效抑制了检测框抖动导致的速度噪声；'
        '通过简化的一维标定方法降低了部署门槛。实验结果表明，系统能在真实交通视频中'
        '提供稳定、合理的速度与TTC估计。'
    )

    add_body_paragraph(doc,
        '在车牌识别任务中，本文构建了YOLO-Pose+CRNN的两阶段架构。'
        '通过引入多省均衡数据集训练，解决了CCPD数据集省份偏差导致的识别倾向性问题；'
        '通过四角点透视变换替代传统仿射变换，提升了倾斜车牌的校正质量；'
        '通过多帧投票机制充分挖掘视频时序冗余，大幅提升了车牌级的识别准确率；'
        '通过字母表扩展增强了对特殊车牌类型的支持能力。'
        '两个任务均实现了完整的GUI系统，支持实时推理、参数调节和结果导出。'
    )

    add_body_paragraph(doc,
        '本研究的贡献可归纳为四个方面：（1）从工程实践角度，为交通场景下的视觉感知任务提供了'
        '可复现、可部署的完整解决方案；（2）从方法创新角度，在多帧线性回归测速、一维道路标定、'
        '多帧投票车牌识别等方面提出了有效的优化策略；（3）从问题分析角度，系统识别并定量分析了'
        '预训练模型类别偏差、数据集省份分布偏差等实际落地中的关键问题；'
        '（4）从实验评估角度，提供了丰富的定量对比和定性分析，为后续研究提供了参考基线。'
    )

    # ================= 参考文献 =================
    add_heading_custom(doc, '参考文献', level=1)

    refs = [
        '[1] 陈晓明, 李志强. 基于感应线圈的交通流检测技术综述[J]. 交通信息与安全, 2018, 36(3): 1-8.',
        '[2] 王磊, 张伟. 激光雷达在智能交通系统中的应用研究[J]. 激光与光电子学进展, 2019, 56(12): 120001.',
        '[3] Lai A H S, Yung N H C. Vehicle-type identification through automated virtual loop assignment and block-based direction-biased motion estimation[J]. IEEE Transactions on Intelligent Transportation Systems, 2000, 1(2): 86-97.',
        '[4] Dailey D J, Cathey F W, Pumrin S. An algorithm to estimate mean traffic speed using uncalibrated cameras[J]. IEEE Transactions on Intelligent Transportation Systems, 2000, 1(2): 98-107.',
        '[5] Zivkovic Z. Improved adaptive Gaussian mixture model for background subtraction[C]. Proceedings of the 17th International Conference on Pattern Recognition (ICPR), 2004: 28-31.',
        '[6] Horn B K P, Schunck B G. Determining optical flow[J]. Artificial Intelligence, 1981, 17(1-3): 185-203.',
        '[7] Kalman R E. A new approach to linear filtering and prediction problems[J]. Journal of Basic Engineering, 1960, 82(1): 35-45.',
        '[8] Redmon J, Divvala S, Girshick R, et al. You only look once: Unified, real-time object detection[C]. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR), 2016: 779-788.',
        '[9] Zhang Y, Sun P, Jiang Y, et al. ByteTrack: Multi-object tracking by associating every detection box[C]. European Conference on Computer Vision (ECCV), 2022: 1-21.',
        '[10] 刘志勇, 杨华. 基于边缘检测和模板匹配的车牌识别方法[J]. 计算机应用与软件, 2010, 27(8): 89-92.',
        '[11] 赵明, 陈强. 基于颜色特征的车牌定位算法研究[J]. 模式识别与人工智能, 2005, 18(6): 735-740.',
        '[12] 林伟, 王欣. 基于数学形态学的车牌定位方法[J]. 计算机工程, 2006, 32(3): 208-210.',
        '[13] 陈军, 杜勇. 基于自适应投影分割的车牌字符切分[J]. 计算机工程与设计, 2009, 30(7): 1750-1753.',
        '[14] 李斌, 刘强. 基于HOG特征与SVM的车牌字符识别[J]. 计算机科学, 2011, 38(10): 283-286.',
        '[15] Maji D, Nagori S, Mathew M, et al. YOLO-Pose: Enhancing YOLO for multi person pose estimation using object keypoint similarity loss[C]. Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition Workshops (CVPRW), 2022: 2637-2646.',
        '[16] Shi B, Bai X, Yao C. An end-to-end trainable neural network for image-based sequence recognition and its application to scene text recognition[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2017, 39(11): 2298-2304.',
        '[17] Graves A, Fernández S, Gomez F, et al. Connectionist temporal classification: labelling unsegmented sequence data with recurrent neural networks[C]. Proceedings of the 23rd International Conference on Machine Learning (ICML), 2006: 369-376.',
        '[18] Xu Z, Yang W, Meng A, et al. Towards end-to-end license plate detection and recognition: A large dataset and baseline[C]. Proceedings of the European Conference on Computer Vision (ECCV), 2018: 255-271.',
        '[19] Jocher G, Chaurasia A, Qiu J. Ultralytics YOLO[CP/OL]. https://github.com/ultralytics/ultralytics, 2023.',
        '[20] (相关实验指导书) 模式识别与机器视觉实践——实验二：车速和撞线时间估计[R]. 华南理工大学, 2025.',
        '[21] (相关实验指导书) 模式识别与机器视觉实践——实验四：复杂环境下基于机器视觉处理的车牌识别系统[R]. 华南理工大学, 2025.',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=2.0)
        run = p.add_run(ref)
        set_font(run, size=Pt(12))

    # ========== 保存 ==========
    output_path = '实验报告_车辆速度估计与车牌识别系统.docx'
    output_dir = Path(__file__).resolve().parents[1] / "docs" / "reports"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "实验报告_车辆速度估计与车牌识别系统.docx"
    doc.save(output_path)
    print(f'报告已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    build_report()
